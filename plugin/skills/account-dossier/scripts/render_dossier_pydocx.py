#!/usr/bin/env python3
"""
render_dossier_pydocx.py — python-docx renderer for the account-dossier skill.

Companion to scripts/build_dossier.js (docx-js). Reads the SAME declarative JSON
spec (see references/document-structure.md) and produces an equivalent .docx via
python-docx, which is already a pyproject dependency — no npm/node `docx` package
needed. Use this whenever `node build_dossier.js` fails with
`Cannot find module 'docx'` (the node package is not installed in this
environment and `npm install`/`npx` are hard-denied — see
agent/permissions.py:_DANGEROUS_PROGRAMS).

Usage:  uv run python scripts/render_dossier_pydocx.py <spec.json> <out.docx>

Four OOXML gotchas are baked in below (each one cost a from-scratch
rediscovery the first time this was written, throwaway, in a scratchpad):

  1. WIDTH UNITS. `tcW`/`tblW` widths (CT_TblWidth.w) are a plain integer
     number of DXA/twips — NOT an EMU-based python-docx Length. Assigning an
     Inches()/Pt()-derived Length object's raw int to `.w` writes a number
     ~635x too large (EMU-per-twip) and silently corrupts the layout. `gridCol.w`
     (CT_TblGridCol.w) is the opposite: it DOES want a Length (ST_TwipsMeasure),
     so `Twips(n)` there, plain `n` on `tcW`/`tblW`. See `DXA()` / `_set_cell_width()`.
  2. ELEMENT ORDER. `CT_TcPr`, `CT_PPr`, and `CT_TblPr` are strict-order XML
     sequences — appending a new child at the end is invalid if the schema
     expects it earlier (e.g. `tcBorders` must precede `tcMar`). `_insert_ordered()`
     inserts each raw element at its correct schema position instead of blindly
     appending.
  3. BORDER SIDE NAMES. This environment's docx validator (the one bundled with
     the anthropic-skills `docx` skill) accepts "start"/"end" for
     `tcBorders`/`tblBorders` sides and REJECTS the "left"/"right" compatibility
     aliases there — even though "left"/"right" is required and correct for
     `tcMar` (cell margins). Easy to get backwards; `_BORDER_SIDES` codifies it.
  4. ZOOM PERCENT. python-docx's own `Document()` template ships
     `<w:zoom w:val="bestFit"/>` with no `w:percent` attribute, which the same
     validator flags as invalid. `_fix_zoom_element()` patches it post-creation.
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from PIL import Image

# ── brand defaults (overridable per spec, mirrors build_dossier.js) ────────────
DEFAULT_BRAND = {
    "bg": "0A0E27",
    "accent1": "3E7BFA",
    "accent2": "B57BFF",
    "accent3": "22D3EE",
    "textOnDark": "FFFFFF",
    "font": "Calibri",
}

# US Letter content width (12240 − 2×720 margins), in DXA/twips — matches
# build_dossier.js's CONTENT_W exactly.
MARGIN_DXA = 720
PAGE_W_DXA = 12240
PAGE_H_DXA = 15840
CONTENT_W_DXA = PAGE_W_DXA - 2 * MARGIN_DXA

TABLE_BORDER_COLOR = "CCCCCC"

# Gotcha #2: hardcoded strict-order child sequences (ECMA-376 §17.4). Copied
# here rather than read off python-docx's private `_tag_seq` because several
# of these classes `del _tag_seq` at the end of their class body, making it
# unavailable at runtime.
_TCPR_ORDER = [
    "w:cnfStyle",
    "w:tcW",
    "w:gridSpan",
    "w:hMerge",
    "w:vMerge",
    "w:tcBorders",
    "w:shd",
    "w:noWrap",
    "w:tcMar",
    "w:textDirection",
    "w:tcFitText",
    "w:vAlign",
    "w:hideMark",
    "w:headers",
    "w:cellIns",
    "w:cellDel",
    "w:cellMerge",
    "w:tcPrChange",
]
_PPR_ORDER = [
    "w:pStyle",
    "w:keepNext",
    "w:keepLines",
    "w:pageBreakBefore",
    "w:framePr",
    "w:widowControl",
    "w:numPr",
    "w:suppressLineNumbers",
    "w:pBdr",
    "w:shd",
    "w:tabs",
    "w:suppressAutoHyphens",
    "w:kinsoku",
    "w:wordWrap",
    "w:overflowPunct",
    "w:topLinePunct",
    "w:autoSpaceDE",
    "w:autoSpaceDN",
    "w:bidi",
    "w:adjustRightInd",
    "w:snapToGrid",
    "w:spacing",
    "w:ind",
    "w:contextualSpacing",
    "w:mirrorIndents",
    "w:suppressOverlap",
    "w:jc",
    "w:textDirection",
    "w:textAlignment",
    "w:textboxTightWrap",
    "w:outlineLvl",
    "w:divId",
    "w:cnfStyle",
    "w:rPr",
    "w:sectPr",
    "w:pPrChange",
]
_TBLPR_ORDER = [
    "w:tblStyle",
    "w:tblpPr",
    "w:tblOverlap",
    "w:bidiVisual",
    "w:tblStyleRowBandSize",
    "w:tblStyleColBandSize",
    "w:tblW",
    "w:jc",
    "w:tblCellSpacing",
    "w:tblInd",
    "w:tblBorders",
    "w:shd",
    "w:tblLayout",
    "w:tblCellMar",
    "w:tblLook",
    "w:tblCaption",
    "w:tblDescription",
    "w:tblPrChange",
]

# Gotcha #3: border-side element names — "start"/"end", never "left"/"right".
_BORDER_SIDES = ("top", "start", "bottom", "end")


def _insert_ordered(parent_elm, order, child_elm):
    """Insert `child_elm` into `parent_elm` at its correct schema position.

    `order` is the full strict-order tag sequence for `parent_elm`'s type
    (one of _TCPR_ORDER / _PPR_ORDER / _TBLPR_ORDER). Existing children with
    an earlier-or-equal position are left alone; child_elm is placed just
    before the first existing child that must come after it.
    """
    tag = child_elm.tag
    # tag is a Clark-notation {ns}local string; order entries are "w:local".
    local = tag.split("}", 1)[1] if "}" in tag else tag
    idx = order.index(f"w:{local}")
    successors = [
        f"w:{t.split(':', 1)[1]}" if not t.startswith("w:") else t for t in order[idx + 1 :]
    ]
    successor_qns = [qn(t) for t in successors]
    for existing in list(parent_elm):
        if existing.tag in successor_qns:
            existing.addprevious(child_elm)
            return child_elm
    parent_elm.append(child_elm)
    return child_elm


def hx(color):
    return str(color or "").lstrip("#").upper() or "000000"


def hp(size_halfpt):
    """Half-points (docx-js `size`) -> Pt(), so callers can copy JS sizes verbatim."""
    return Pt(size_halfpt / 2)


class Renderer:
    def __init__(self, spec, spec_dir):
        self.spec = spec
        self.spec_dir = spec_dir
        brand = dict(DEFAULT_BRAND)
        brand.update({k: hx(v) for k, v in (spec.get("brand") or {}).items()})
        self.brand = brand
        self.font = brand["font"]
        self.doc = Document()
        self._setup_document()

    # ── document scaffolding ────────────────────────────────────────────────
    def _setup_document(self):
        doc = self.doc
        normal = doc.styles["Normal"]
        normal.font.name = self.font
        normal.font.size = Pt(11)

        section = doc.sections[0]
        section.page_width = Twips(PAGE_W_DXA)
        section.page_height = Twips(PAGE_H_DXA)
        section.top_margin = Twips(MARGIN_DXA)
        section.bottom_margin = Twips(MARGIN_DXA)
        section.left_margin = Twips(MARGIN_DXA)
        section.right_margin = Twips(MARGIN_DXA)

        self._fix_zoom_element()
        self._build_footer(section)

    def _fix_zoom_element(self):
        # Gotcha #4: patch the template's zoom element so it has a percent attr.
        settings = self.doc.settings.element
        zoom = settings.find(qn("w:zoom"))
        if zoom is not None and zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")

    def _build_footer(self, section):
        closing = (
            self.spec.get("closingLine") or "Internal briefing — not for customer distribution."
        )
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{closing}   ·   Page ")
        r.font.name = self.font
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        self._add_page_number_field(p)

    def _add_page_number_field(self, paragraph):
        run = paragraph.add_run()
        run.font.name = self.font
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    # ── low-level OOXML helpers ─────────────────────────────────────────────
    def _set_cell_width(self, cell, dxa):
        """Gotcha #1: raw DXA int on tcW, never a Length-derived number."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            _insert_ordered(tcPr, _TCPR_ORDER, tcW)
        tcW.set(qn("w:w"), str(int(dxa)))
        tcW.set(qn("w:type"), "dxa")

    def _set_table_grid(self, table, column_widths_dxa):
        table.autofit = False
        tblPr = table._tbl.tblPr
        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            _insert_ordered(tblPr, _TBLPR_ORDER, layout)
        layout.set(qn("w:type"), "fixed")
        for i, col in enumerate(table.columns):
            col.width = Twips(column_widths_dxa[i])
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            _insert_ordered(tblPr, _TBLPR_ORDER, tblW)
        tblW.set(qn("w:w"), str(int(sum(column_widths_dxa))))
        tblW.set(qn("w:type"), "dxa")

    def _set_cell_shading(self, cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            _insert_ordered(tcPr, _TCPR_ORDER, shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hx(fill_hex))

    def _set_cell_margins(self, cell, top=80, bottom=80, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is None:
            tcMar = OxmlElement("w:tcMar")
            _insert_ordered(tcPr, _TCPR_ORDER, tcMar)
        for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            el = tcMar.find(qn(f"w:{side}"))
            if el is None:
                el = OxmlElement(f"w:{side}")
                tcMar.append(el)  # tcMar's own children have a fixed order too, but
                # top/left/bottom/right are inserted in that exact fixed order below.
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")
        # tcMar child order is top, left, bottom, right — rebuild in order to be safe.
        order = ["top", "left", "bottom", "right"]
        children = {c.tag.split("}", 1)[1]: c for c in list(tcMar)}
        for c in list(tcMar):
            tcMar.remove(c)
        for side in order:
            if side in children:
                tcMar.append(children[side])

    def _set_cell_borders(self, cell, color=TABLE_BORDER_COLOR, sz=4):
        # Gotcha #3: "start"/"end", never "left"/"right", inside tcBorders.
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            _insert_ordered(tcPr, _TCPR_ORDER, tcBorders)
        for side in _BORDER_SIDES:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), hx(color))
            existing = tcBorders.find(qn(f"w:{side}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(el)

    def _add_heading_rule(self, paragraph, color):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            _insert_ordered(pPr, _PPR_ORDER, pBdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), hx(color))
        pBdr.append(bottom)

    def _add_hyperlink(self, paragraph, url, text, size_halfpt=18):
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self.font)
        rFonts.set(qn("w:hAnsi"), self.font)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size_halfpt))
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        hyperlink.append(r)
        paragraph._p.append(hyperlink)

    # ── text/paragraph helpers ──────────────────────────────────────────────
    def _run(self, container, text, size=22, bold=False, italic=False, color=None):
        r = container.add_run("" if text is None else str(text))
        r.font.name = self.font
        r.font.size = hp(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor.from_string(hx(color))
        return r

    def _paragraph(self, text, size=22, spacing_after=120):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Twips(spacing_after)
        self._run(p, text, size=size)
        return p

    def _bullet(self, cell_or_doc, text, size=20, color=None, italic=False, spacing_after=30):
        p = cell_or_doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Twips(spacing_after)
        self._run(p, text, size=size, color=color, italic=italic)
        return p

    def _new_table(self, rows, cols, column_widths_dxa):
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = None
        self._set_table_grid(table, column_widths_dxa)
        return table

    # ── block renderers (mirror build_dossier.js's switch cases) ───────────
    def render_banner(self, blk):
        table = self._new_table(1, 1, [CONTENT_W_DXA])
        cell = table.rows[0].cells[0]
        self._set_cell_shading(cell, self.brand["bg"])
        self._set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
        self._set_cell_width(cell, CONTENT_W_DXA)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = cell.paragraphs[0]
        self._run(p0, blk.get("account", ""), size=44, bold=True, color=self.brand["textOnDark"])
        if blk.get("subline"):
            p1 = cell.add_paragraph()
            p1.paragraph_format.space_before = Twips(40)
            self._run(p1, blk["subline"], size=24, color=self.brand["textOnDark"])
        if blk.get("meta"):
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Twips(40)
            self._run(p2, blk["meta"], size=18, color=self.brand["accent3"])
        self.doc.add_paragraph().paragraph_format.space_after = Twips(120)

    def render_callout(self, blk):
        lines = blk.get("lines") or [blk.get("text", "")]
        variant = blk.get("variant")
        left_color = self.brand["accent2"] if variant == "angle" else self.brand["accent3"]
        table = self._new_table(1, 1, [CONTENT_W_DXA])
        cell = table.rows[0].cells[0]
        self._set_cell_shading(cell, "F0F4FF")
        self._set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
        self._set_cell_width(cell, CONTENT_W_DXA)
        # Left accent bar: single-sided border, not the full box the guardrails warn against.
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        _insert_ordered(tcPr, _TCPR_ORDER, tcBorders)
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "single")
        start.set(qn("w:sz"), "24")
        start.set(qn("w:space"), "0")
        start.set(qn("w:color"), hx(left_color))
        tcBorders.append(start)
        for i, ln in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Twips(0 if i == len(lines) - 1 else 60)
            self._run(p, ln, size=22, bold=(variant == "angle"))
        self.doc.add_paragraph().paragraph_format.space_after = Twips(120)

    def render_heading(self, blk):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Twips(260)
        p.paragraph_format.space_after = Twips(120)
        self._add_heading_rule(p, self.brand["accent1"])
        self._run(
            p, str(blk.get("text") or "").upper(), size=26, bold=True, color=self.brand["accent1"]
        )

    def render_paragraph(self, blk):
        self._paragraph(blk.get("text", ""), size=22, spacing_after=120)

    def render_facts_table(self, blk):
        rows = blk.get("rows") or []
        if not rows:
            return
        w0 = round(CONTENT_W_DXA * 0.32)
        w1 = CONTENT_W_DXA - w0
        table = self._new_table(len(rows), 2, [w0, w1])
        for r, pair in enumerate(rows):
            label_cell, value_cell = table.rows[r].cells
            self._set_cell_shading(label_cell, "EEF2FB")
            for cell, dxa in ((label_cell, w0), (value_cell, w1)):
                self._set_cell_margins(cell)
                self._set_cell_borders(cell)
                self._set_cell_width(cell, dxa)
            self._run(
                label_cell.paragraphs[0], pair[0] if len(pair) > 0 else "", size=20, bold=True
            )
            self._run(value_cell.paragraphs[0], pair[1] if len(pair) > 1 else "", size=20)
        self.doc.add_paragraph().paragraph_format.space_after = Twips(120)

    def render_table(self, blk):
        header = blk.get("header") or []
        body = blk.get("rows") or []
        cols = len(header) or (len(body[0]) if body else 1)
        if cols == 0:
            return
        w = CONTENT_W_DXA // cols
        widths = [w] * cols
        total_rows = (1 if header else 0) + len(body)
        if total_rows == 0:
            return
        table = self._new_table(total_rows, cols, widths)
        r_idx = 0
        if header:
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                self._set_cell_shading(cell, self.brand["accent1"])
                self._set_cell_margins(cell)
                self._set_cell_borders(cell)
                self._set_cell_width(cell, w)
                self._run(
                    cell.paragraphs[0], text, size=20, bold=True, color=self.brand["textOnDark"]
                )
            r_idx = 1
        for row in body:
            for c, text in enumerate(row):
                cell = table.rows[r_idx].cells[c]
                self._set_cell_margins(cell)
                self._set_cell_borders(cell)
                self._set_cell_width(cell, w)
                self._run(cell.paragraphs[0], text, size=20)
            r_idx += 1
        self.doc.add_paragraph().paragraph_format.space_after = Twips(120)

    def render_two_col(self, blk):
        half = CONTENT_W_DXA // 2
        table = self._new_table(2, 2, [half, half])
        header_titles = (blk.get("leftTitle") or "DO", blk.get("rightTitle") or "DON'T")
        header_fills = (self.brand["accent1"], self.brand["accent2"])
        for c in range(2):
            cell = table.rows[0].cells[c]
            self._set_cell_shading(cell, header_fills[c])
            self._set_cell_margins(cell)
            self._set_cell_borders(cell)
            self._set_cell_width(cell, half)
            self._run(
                cell.paragraphs[0],
                header_titles[c],
                size=20,
                bold=True,
                color=self.brand["textOnDark"],
            )
        body_items = (blk.get("left") or [], blk.get("right") or [])
        for c in range(2):
            cell = table.rows[1].cells[c]
            self._set_cell_margins(cell)
            self._set_cell_borders(cell)
            self._set_cell_width(cell, half)
            items = body_items[c]
            if not items:
                self._run(cell.paragraphs[0], "", size=20)
                continue
            first = True
            for item in items:
                p = cell.paragraphs[0] if first else cell.add_paragraph(style="List Bullet")
                if first:
                    p.style = self.doc.styles["List Bullet"]
                    first = False
                p.paragraph_format.space_after = Twips(40)
                self._run(p, item, size=20)
        self.doc.add_paragraph().paragraph_format.space_after = Twips(120)

    def render_questions(self, blk):
        for g in blk.get("groups") or []:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Twips(80)
            p.paragraph_format.space_after = Twips(40)
            self._run(p, g.get("title", ""), size=22, bold=True, color=self.brand["accent1"])
            for q in g.get("items") or []:
                self._bullet(self.doc, q, size=20, spacing_after=30)

    def render_image(self, blk):
        rel_path = blk.get("path")
        if not rel_path:
            return
        abs_path = rel_path if os.path.isabs(rel_path) else os.path.join(self.spec_dir, rel_path)
        if not os.path.exists(abs_path):
            print(f"image not found, skipping: {rel_path}", file=sys.stderr)
            return
        width_in = blk.get("widthIn") or 6
        height_in = blk.get("heightIn")
        if not height_in:
            # Preserve the image's own aspect ratio rather than guessing a
            # fixed multiplier — a wrong guess stretches/squashes the image
            # and can silently push the doc past its page-count cap.
            try:
                with Image.open(abs_path) as img:
                    px_w, px_h = img.size
                height_in = width_in * (px_h / px_w)
            except Exception as exc:
                print(
                    f"could not read image size, falling back to 0.5 ratio: {exc}", file=sys.stderr
                )
                height_in = width_in * 0.5
        self.doc.add_picture(
            abs_path, width=Twips(int(width_in * 1440)), height=Twips(int(height_in * 1440))
        )
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.paragraphs[-1].paragraph_format.space_before = Twips(80)
        self.doc.paragraphs[-1].paragraph_format.space_after = Twips(120)

    def render_sources(self, blk):
        for s in blk.get("external") or []:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Twips(30)
            if s.get("url"):
                self._add_hyperlink(p, s["url"], s.get("text") or s["url"], size_halfpt=18)
            else:
                self._run(p, s.get("text", ""), size=18)
        for t in blk.get("internal") or []:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Twips(30)
            self._run(p, t, size=18, italic=True)

    def render_spacer(self, _blk):
        self.doc.add_paragraph().paragraph_format.space_after = Twips(120)

    def render_pagebreak(self, _blk):
        self.doc.add_page_break()

    # ── entry point ──────────────────────────────────────────────────────────
    def render(self):
        dispatch = {
            "banner": self.render_banner,
            "callout": self.render_callout,
            "heading": self.render_heading,
            "paragraph": self.render_paragraph,
            "facts_table": self.render_facts_table,
            "table": self.render_table,
            "two_col": self.render_two_col,
            "questions": self.render_questions,
            "image": self.render_image,
            "sources": self.render_sources,
            "spacer": self.render_spacer,
            "pagebreak": self.render_pagebreak,
        }
        blocks = self.spec.get("blocks") or []
        count = 0
        for blk in blocks:
            if not isinstance(blk, dict):
                continue
            fn = dispatch.get(blk.get("type"))
            if fn is None:
                print(f"unknown block type, skipping: {blk.get('type')}", file=sys.stderr)
                continue
            fn(blk)
            count += 1
        if not blocks:
            self._paragraph("(empty dossier spec)")
        return count


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("spec_path")
    parser.add_argument("out_path")
    args = parser.parse_args()

    with open(args.spec_path, encoding="utf-8") as f:
        spec = json.load(f)

    renderer = Renderer(spec, os.path.dirname(os.path.abspath(args.spec_path)))
    count = renderer.render()
    renderer.doc.save(args.out_path)
    size = os.path.getsize(args.out_path)
    print(f"wrote {args.out_path} ({size} bytes, {count} blocks)")


if __name__ == "__main__":
    main()
